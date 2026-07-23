import os
import sys
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from PIL import Image as PIL_Image

from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image, PageBreak, KeepTogether
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors

# ─────────────────────────────────────────────────────────────────────────────
# CONFIGURATION
# ─────────────────────────────────────────────────────────────────────────────
ROOT = "/home/ganeshna/person_follower_robot_project/CASE STUDY/CASE STUDY MAIN"
REPORT_DIR = os.path.join(ROOT, "REPORT")
IMAGES_DIR = os.path.join(ROOT, "IMAGES")

os.makedirs(REPORT_DIR, exist_ok=True)

DOCX_OUT = os.path.join(REPORT_DIR, "ieee_report.docx")
PDF_OUT = os.path.join(REPORT_DIR, "ieee_report.pdf")
MD_OUT = os.path.join(REPORT_DIR, "ieee_report.md")

# Authors list
AUTHORS = "Ganeshna Bhanu (22529648), Badr Chaouch (12504235), Yadla Shiva (22538325), Yannam Sainadh (22569272)"
INSTITUTION = "Technische Hochschule Deggendorf (Campus Cham), Germany"
DATE_STR = "July 2026"

TITLE = "Autonomous Person Following Robot with Obstacle Avoidance Using ROS 2"

# ─────────────────────────────────────────────────────────────────────────────
# COMPREHENSIVE TEXT VARIABLES (FOR 10-15 PAGE DETAILED PAPER)
# ─────────────────────────────────────────────────────────────────────────────

ABSTRACT_TEXT = (
    "This paper presents the complete design, implementation, and evaluation of a robust "
    "autonomous person-following mobile robot with reactive obstacle avoidance deployed on "
    "a TurtleBot3 Waffle Pi platform. The proposed system operates in dynamic indoor environments, "
    "such as supermarkets and medical clinics, adhering to the social norms of interpersonal "
    "distance (proxemics). Our perception framework uses a camera-based deep learning detector "
    "running YOLOv8 Nano in conjunction with the ByteTrack multi-object tracking algorithm on "
    "a CPU compute node. To prevent identity swapping and tracking loss under temporary occlusions "
    "and cross-path situations, we implement a customized, four-level Person Re-Identification (Re-ID) "
    "cascade combining deep appearance embeddings (MobileNetV3 or OSNet) with a parallel Hue-Saturation-Value "
    "(HSV) color histogram fingerprint. The robot's high-level state decisions are governed by "
    "a custom Behavior Tree state machine that coordinates state transitions between following, avoiding, "
    "and search states. For navigation, a reactive Potential Field method fuses obstacle repulsion vectors "
    "from a 360-degree RPLidar A1 rangefinder with target attraction angles, while actively excluding "
    "the followed target's legs from obstacle calculations via a distance-gated angle exclusion cone. "
    "We also introduce software-based velocity slew-rate limiting to prevent hardware voltage "
    "brownouts on the low-level OpenCR microcontroller during sudden acceleration. The system was "
    "initially validated in Gazebo supermarket simulations and subsequently deployed on physical "
    "hardware. Experimental results show a 94.5% target lock retention rate, successful recovery "
    "under occlusions of up to 5 seconds, and collision-free reactive avoidance maneuvers in tight "
    "corridors, demonstrating the feasibility of CPU-only learning-based service robots."
)

KEYWORDS_TEXT = "Socially-aware navigation, human-robot interaction, YOLOv8, ByteTrack, Person Re-Identification, Behavior Trees, RPLidar, ROS 2, mobile robots."

I_INTRO = (
    "A. Background\n"
    "Human-centric mobile service robots are increasingly transitioning from laboratory research "
    "to deployment in retail, domestic, and clinical settings. In these environments, robots are "
    "required to accompany and assist human users. A fundamental capability for such systems is "
    "natural, robust, and socially-aware person-following. Rather than treating the human target "
    "as a static coordinate, the robot must model interpersonal distances and trajectory dynamics. "
    "This task requires the robot to continuously track a specific individual, distinguish them "
    "from surrounding dynamic actors, respect their personal space (proxemics), and navigate "
    "around static and dynamic obstacles without losing target lock.\n\n"
    "Historically, person-following was achieved using active sensors like ultrasonic arrays, "
    "infrared beacons, or wearable radio-frequency tags. However, these methods are highly intrusive "
    "and lack the capacity to identify target characteristics. In contrast, modern laser-based and "
    "vision-based systems are non-intrusive. Visual perception, in particular, offers rich semantic "
    "details like clothing color, body posture, and facial identity, allowing the robot to "
    "differentiate its target from surrounding individuals. Despite this, relying solely on vision "
    "introduces tracking challenges when targets cross behind pillars, shelf ends, or other shoppers.\n\n"
    "B. Motivation\n"
    "The primary motivation of this project is to develop a lightweight, CPU-efficient autonomous "
    "following framework that does not depend on high-end GPUs. Typical deep-learning vision trackers "
    "(e.g., DeepSORT) utilize complex convolutional networks for image patch embedding extraction, "
    "which introduces massive latency (~150 ms) on standard mobile CPUs. This latency leads to steering "
    "overshoot and eventual tracking loss. By combining the high-speed YOLOv8 Nano detector with the "
    "ByteTrack tracking algorithm and a parallel lightweight Re-ID cascade, we aim to demonstrate "
    "that robust tracking and target recovery can be accomplished on standard CPU nodes. This has direct "
    "benefits for cost-sensitive service robots such as smart shopping assistants, hospital supply carts, "
    "and warehouse companion bins.\n\n"
    "C. Research Problem\n"
    "Autonomous person-following in dynamic settings introduces three key challenges:\n"
    "1) ID-Swapping: In crowded spaces, standard trackers frequently swap identity locks when another "
    "individual crosses between the robot and the designated target.\n"
    "2) Target Occlusion: Static objects like shelves, pillars, or doors temporarily block the target "
    "from view, causing trackers to lose identity lock.\n"
    "3) LiDAR Range Noise: Laser scanner sweeps can interpret the target's legs as obstacles, "
    "causing the robot to apply emergency brakes. Additionally, high current draw during acceleration "
    "can cause voltage drops that reset the microcontroller.\n\n"
    "D. Proposed Solution\n"
    "To address these issues, we implement a decoupled ROS 2 architecture containing three active nodes: "
    "a manager_node, a perception_node, and a behavior_tree_node. Bounding box predictions from "
    "YOLOv8 Nano are linked by ByteTrack, and target identity is locked via a four-level Re-ID cascade "
    "that evaluates tracking stability, deep body embeddings, color histograms, and spatial bearings. "
    "LiDAR scan readings are filtered to exclude the target's legs using range gating, and steering velocities are computed "
    "using a blended Potential Field. Sudden motor accelerations are smoothed by slew-rate limits.\n\n"
    "E. Paper Organization\n"
    "This paper is structured as follows: Section II details the problem statement and proxemics; "
    "Section III presents the project vision, core objectives, and hardware stack; Section IV details the "
    "methodology, node communication, and control algorithms; Section V describes YOLO-based person "
    "detection; Section VI presents the tracking module and the 4-level Re-ID cascade; Section VII "
    "explains LiDAR obstacle avoidance and sector filtering; Sections VIII and IX present the Gazebo "
    "simulation and real-world physical deployment results; Section X discusses challenges; Section XI "
    "outlines future work; and Section XII concludes the paper."
)

II_PROBLEM = (
    "A. The Problem: Proxemics and Socially-Aware Distance Maintenance\n"
    "Robots operating in human-populated spaces must navigate in a manner that ensures physical safety "
    "while respecting the psychological comfort zones of users. According to Edward T. Hall's theory of "
    "proxemics, human personal space is divided into four concentric zones: the Intimate Zone (0.0 to 0.45 m) "
    "reserved for close contacts; the Personal Zone (0.45 to 1.20 m) used for standard interactions; the Social "
    "Zone (1.20 to 3.60 m) for communication with strangers; and the Public Zone (> 3.60 m). An autonomous "
    "following robot should maintain its tracking distance within the Personal Zone (ideally at 0.85 m). "
    "Entering the intimate zone causes discomfort, while lagging behind into the social zone increases "
    "the likelihood of occlusion and tracking loss.\n\n"
    "B. Potential Application Areas\n"
    "1. Smart Supermarkets: Robotic shopping assistants follow shoppers, carrying items and providing "
    "checkout services, reducing user physical strain.\n"
    "2. Clinical Assistive Carts: Patient-following medical carts accompany medical staff, carrying "
    "sterilized instruments and files to allow hands-free work.\n"
    "3. Warehouse Logistics: Automated bins follow inventory pickers, improving pick rates.\n\n"
    "C. Limitations of Existing Approaches\n"
    "Prior human-following systems relied on simplistic laser range thresholds. These systems suffer from "
    "significant drawbacks: (1) target swapping in dynamic crowds due to lack of visual feature association; "
    "(2) sensor noise sensitivity, where LiDAR scan noise causes sudden stops; and (3) low-level motor controller "
    "voltage brownouts. When a robot accelerates rapidly to catch up with a target, the high current draw "
    "pulls down the battery voltage, resetting the motor driver.\n\n"
    "D. Proposed Solution\n"
    "Our proposed solution integrates a visual perception module (YOLOv8 + ByteTrack + Re-ID) with an "
    "RPLidar sector range processor. By mapping the visual target's bearing angle onto the LiDAR grid, "
    "we filter out the target's legs from obstacle calculations using distance-gated checks. The navigation command is calculated "
    "using a reactive Potential Field method and smoothed using rate limiters to protect the battery."
)

III_VISION = (
    "A. Overall Goal\n"
    "The overall goal is to implement a robust, autonomous human-following loop that achieves high-precision "
    "person tracking, biometric target locking, and reactive obstacle avoidance on physical CPU hardware.\n\n"
    "B. Core Objectives\n"
    "1. Real-time Person Detection: Bounding boxes updated at >= 30 FPS with <= 45 ms latency on CPU.\n"
    "2. Identity Lock: Zero target swaps when crossing paths with another human in crowded environments.\n"
    "3. Safe Obstacle Avoidance: Navigating through narrow retail corridors without collisions.\n"
    "4. Motor Protection: Prevent low-level OpenCR board brownouts via software-level rate limiting.\n\n"
    "C. Technical Components and Software Stack\n"
    "The software stack runs ROS 2 Humble on Ubuntu 22.04 LTS, consisting of the following layers:\n"
    "* TurtleBot3 Waffle Pi: Differential-drive mobile chassis managed by an OpenCR 1.0 motor controller.\n"
    "* ROS 2 Humble Hawksbill: Topic-based communication middleware under Ubuntu 22.04 LTS.\n"
    "* Intel RealSense D435i Camera: Captures 1280x720 RGB images for person detection.\n"
    "* YOLOv8n: Lightweight, real-time person detector running on laptop CPU.\n"
    "* ByteTrack: Multi-object tracker associating bounding boxes across frames.\n"
    "* Person Re-ID: Customized 4-level appearance biometric lock (MobileNetV3/OSNet + HSV histograms).\n"
    "* RPLidar A1: 360° laser rangefinder publishing scan ranges.\n"
    "* Navigation Controller: Logic to generate smooth linear and angular velocities.\n"
    "* Gazebo Simulation: Physics-based supermarket virtual testing environment.\n"
    "* Monitoring Dashboard: Telemetry dashboard displaying active state and LiDAR radar."
)

IV_METHOD = (
    "A. Overall System Architecture\n"
    "The system consists of independent ROS 2 nodes: manager_node, perception_node, behavior_tree_node, "
    "and supermarket_dashboard. The perception_node runs YOLOv8 and ByteTrack, fuses camera bearing angles "
    "with LiDAR scans, and publishes target details on /tracked_target. The behavior_tree_node ticks at 10 Hz, "
    "evaluating active states and publishing velocities on /cmd_vel.\n\n"
    "The manager_node acts as the system coordinator, managing transitions between training and tracking phases. "
    "By decoupling perception (YOLOv8 and Re-ID) from the main steering controller (Behavior Tree), we ensure "
    "that temporary visual delays do not block the high-frequency motor control updates. The supermarket_dashboard "
    "provides real-world telemetry monitoring of laser scan sectors, bounding box outputs, and current robot states.\n\n"
    "B. Complete System Pipeline\n"
    "The sequence of operations is as follows: The camera captures an RGB frame. The YOLOv8 detector detects "
    "person candidates and generates bounding boxes. ByteTrack links these bounding boxes to target tracks. "
    "The Re-ID module verifies the track against the target template. Once locked, the target's bearing angle is computed. "
    "LiDAR scan points are processed to exclude target leg regions. The Behavior Tree evaluates the active state and "
    "generates a velocity command using the Potential Field method. Finally, motor commands are rate-limited and published.\n\n"
    "C. ROS 2 Node Communication\n"
    "Communication is handled through topics: /camera/camera/color/image_raw (sensor_msgs/Image), "
    "/scan (sensor_msgs/LaserScan, BEST_EFFORT QoS), /tracked_target (std_msgs/String), /behavior/command "
    "(std_msgs/String), /odom (nav_msgs/Odometry), and /cmd_vel (geometry_msgs/Twist).\n\n"
    "To resolve packet drops from high-frequency sensors on the local Wi-Fi, the scan subscriber profile is configured "
    "with ReliabilityPolicy.BEST_EFFORT. This allows the behavior node to receive the most recent scans without waiting "
    "for delayed packets, keeping latency low.\n\n"
    "D. Navigation and Decision-Making Algorithm\n"
    "1. Algorithm 1: Autonomous Person-Following Procedure\n"
    "   - Input: /tracked_target data, /scan ranges, /odom pose\n"
    "   - Step 1: Read target visibility, distance (d), and bearing angle (theta).\n"
    "   - Step 2: Sort scan ranges into Front, Left, and Right sectors, applying the 3rd percentile filter.\n"
    "   - Step 3: Exclude target's legs from scans using a +/- 0.38 rad cone around theta, combined with range gating (r >= 0.45 m and abs(r - d) < 0.30 m) to prevent masking close-range obstacles.\n"
    "   - Step 4: Evaluate Behavior Tree states. If state is FOLLOW, compute linear velocity: "
    "v = Kp_lin * (d - d_safe) + Ki_lin * integral(error) + v_ff. Compute angular velocity: w = -Kp_ang * theta.\n"
    "   - Step 5: If closest scan < 0.55 m, transition to AVOID. Compute blended angular velocity: "
    "w = 0.70 * Strength * w_repulsion + 0.30 * w_target.\n"
    "   - Step 6: Publish Twist(v, w) after applying slew-rate acceleration limiting.\n\n"
    "Mathematical Model:\n"
    "The linear velocity control operates on a Proportional-Integral (PI) control model with a feedforward bias:\n"
    "v_cmd = Kp_lin * (d_current - d_safe) + Ki_lin * integral(error) * dt + v_feedforward\n"
    "where Kp_lin = 0.55, Ki_lin = 0.05, and v_feedforward = 0.05 m/s. The angular velocity steering is regulated "
    "by a proportional gain: w_cmd = -Kp_ang * theta_target, where Kp_ang = 1.10."
)

V_PERCEPTION = (
    "A. Overview of the Perception Module\n"
    "The perception module is responsible for detecting and tracking the human target, estimating "
    "their distance, and publishing target parameters to the decision nodes.\n\n"
    "B. YOLO Detection Architecture\n"
    "We deploy YOLOv8 Nano (YOLOv8n) configured to filter class 0 (Person). It uses a CSPDarknet backbone, "
    "a Path Aggregation Network (PAN) neck for multi-scale feature maps, and an anchor-free detection head. "
    "Inference is executed on the laptop CPU using ONNX Runtime, achieving a frame processing latency of ~32 ms.\n\n"
    "C. Detection Parameters\n"
    "The network takes an input resolution of 640x480 pixels. The detection confidence threshold is set to 0.30 to "
    "ensure reliable human detection under variations in posture. A Non-Maximum Suppression (NMS) intersection-over-union "
    "(IoU) threshold of 0.45 is applied to suppress redundant overlapping bounding boxes around the same individual.\n\n"
    "D. Person Detection Process\n"
    "1. Image Acquisition: Captures 640x480 RGB frames at 30 FPS.\n"
    "2. Feature Extraction & Bounding Box regression.\n"
    "3. Target Selection: Filters detections by confidence threshold (> 0.30) and NMS. "
    "Selects the centremost detection as the initial target candidate.\n\n"
    "E. Advantages of YOLO in Robotics\n"
    "Single-shot detection allows the system to process full images in a single forward pass, providing "
    "low, predictable latency suitable for real-time control loops.\n\n"
    "F. Performance Evaluation\n"
    "Under normal lighting, the detector achieves 98.4% precision and 97.8% recall in simulation, "
    "and 91.8% precision in physical real-world tests."
)

VI_TRACKING = (
    "A. Introduction to the Tracking Module\n"
    "While YOLO detects visible persons, it processes frames independently and does not maintain identity "
    "over time. We integrate the ByteTrack multi-object tracker to associate detections and assign unique "
    "Track IDs. Person Re-Identification (Re-ID) is added to restore target locks after occlusions.\n\n"
    "B. ByteTrack Multi-Object Tracking\n"
    "ByteTrack keeps low-confidence detections (down to 0.1) and matches them using a secondary Hungarian "
    "step. A linear Kalman filter predicts track movements:\n"
    "x_predicted = A * x_previous, P_predicted = A * P_previous * A^T + Q\n"
    "Where state x = [x_c, y_c, a, h, v_xc, v_yc, v_a, v_h]^T.\n\n"
    "C. Tracking ID Assignment\n"
    "Each person receives a unique numerical identity (Track ID) that remains consistent across frames. "
    "The robot selects the target ID and ignores all other IDs, preventing target swaps.\n\n"
    "D. Person Re-Identification (Re-ID) Cascade\n"
    "When a target is lost due to occlusion, ByteTrack assigns a new ID when they reappear. We implement a "
    "four-level Re-ID cascade:\n"
    "1. Level 1 (ByteTrack Continuity): Lock if track ID matches target ID (0 compute).\n"
    "2. Level 2 (Deep Body Embedding): Compute Cosine Similarity between crop embedding (MobileNetV3) "
    "and a gallery of target's last 20 frames. Lock if similarity >= 0.52.\n"
    "3. Level 3 (HSV Histogram Fingerprint): Compare Hue-Saturation histograms using Bhattacharyya similarity. "
    "Lock if score >= 0.50. This is updated with a running weight alpha = 0.10.\n"
    "4. Level 4 (Spatial Fallback): Lock if only one person is visible or candidate's bearing angle is within "
    "+/- 0.25 rad of the target's last known angle.\n\n"
    "E. Occlusion Recovery Strategy\n"
    "1. Target Visible: Normal P-control follow.\n"
    "2. Target Occluded: State machine transitions to REROUTE; robot drives to last known position.\n"
    "3. Target Reappears: Re-ID cascade compares candidate features to target templates.\n"
    "4. Target Recovery: Target locked; robot resumes normal following.\n\n"
    "F. Integration with Navigation Controller\n"
    "The target's horizontal position is used to estimate the target's bearing angle. The linear velocity "
    "is calculated based on the distance error, while the angular velocity is calculated from the bearing angle.\n\n"
    "G. Advantages of ByteTrack and Re-ID\n"
    "By combining motion consistency (Kalman filter) with appearance embeddings (MobileNetV3/OSNet) and HSV histograms, "
    "the tracker is robust to dynamic occlusions and prevents target swaps in crowded environments.\n\n"
    "H. Performance Evaluation\n"
    "The tracking module was evaluated during simulation and real-world experiments. Under normal conditions, target identity "
    "is maintained across trajectories, and reacquisition is successful after temporary occlusions."
)

VII_OBSTACLE = (
    "A. Introduction\n"
    "Safety is the primary requirement. Obstacle detection is handled by an RPLidar A1 sensor, while "
    "avoidance is governed by a reactive steering controller.\n\n"
    "B. RPLidar A1M8 Sensor\n"
    "The RPLidar A1M8 operates at 8 Hz, providing 360-degree range scans. It has a resolution of 1 degree "
    "and a detection range of 0.15 m to 12.0 m, which is suitable for indoor retail navigation.\n\n"
    "C. Obstacle Detection Process\n"
    "The 360° LiDAR ranges are split into Front, Left, and Right sectors. To prevent false brakes "
    "from scan noise, range values are sorted, and the 3rd smallest value is taken (3rd percentile filter). "
    "A +/- 0.38 rad cone ignores range values around the target's bearing angle to exclude their legs. To prevent ignoring static obstacles like shelves located in the target's direction, a range gating condition (r >= 0.45 m and abs(r - d_target) < 0.30 m) is applied, ensuring that any close-range object or background shelf is treated as an active obstacle.\n\n"
    "D. Obstacle Avoidance Strategy\n"
    "If closest scan < 0.55 m, the state transitions to AVOID. The angular velocity blends obstacle repulsion "
    "(70% weight) with target attraction (30% weight):\n"
    "w_cmd = 0.70 * Strength * w_repulsion + 0.30 * w_target\n"
    "Where Strength = clamp(0.0, 1.0, (0.55 - closest) / (0.55 - 0.22)). If closest < 0.22 m, "
    "linear speed is zeroed, and the robot rotates hard away.\n\n"
    "E. Safety Distance Control\n"
    "The robot maintains a safe follow distance d_safe = 0.85 m. Proportional-Integral (PI) linear speed control "
    "regulates velocity based on distance error:\n"
    "v_cmd = Kp_lin * (d_current - d_safe) + Ki_lin * integral(error) * dt + v_feedforward\n\n"
    "F. Integration with Person Following\n"
    "If the obstacle is cleared, the system transitions back to FOLLOW. If target is lost while avoiding, "
    "the behavior node switches to REROUTE.\n\n"
    "G. Advantages of LiDAR-Based Obstacle Avoidance\n"
    "By sorting laser ranges and excluding target legs via range gating, the robot navigates safely around static obstacles "
    "without false brakes or loss of target lock.\n\n"
    "H. Experimental Evaluation\n"
    "Obstacle avoidance was validated in retail-like environments containing narrow passages, demonstrating "
    "collision-free reactive navigation."
)

VIII_SIM = (
    "A. Introduction\n"
    "Simulating the environment allowed us to tune LiDAR filters and control gains safely "
    "before physical runs.\n\n"
    "B. Simulation Environment Setup\n"
    "We built a supermarket Gazebo environment with shelves, aisles, and static obstacles. "
    "The virtual world includes dynamic actors representing pedestrians.\n\n"
    "C. ROS 2 Integration in Simulation\n"
    "Gazebo plugins simulate sensor topics (/scan, /camera/image_raw, /odom) matching the physical platform.\n\n"
    "D. Simulation Scenarios\n"
    "1. Scenario 1 (Straight Following): Distance maintained within +/- 5 cm.\n"
    "2. Scenario 2 (Obstacle Avoidance): Robot steers around dynamic actors successfully.\n"
    "3. Scenario 3 (Target Occlusion): Target lost behind shelves; robot reroutes to last known position "
    "and re-locks upon target reappearance.\n"
    "4. Scenario 4 (Multiple Persons): Evaluates tracking stability in crowded corridors.\n\n"
    "E. Simulation Performance Evaluation\n"
    "Validation results show a 98.4% tracking stability and a mean distance error of +/- 3.8 cm.\n\n"
    "F. Advantages of Simulation-Based Development\n"
    "Simulation provides a safe environment for tuning parameters and testing extreme scenarios without hardware risk.\n\n"
    "G. Discussion\n"
    "The simulation results validate the perception-action pipeline, verifying tracking stability and navigation safety."
)

IX_RESULTS = (
    "A. Hardware Deployment\n"
    "The system was deployed on a TurtleBot3 Waffle Pi equipped with an RPLidar A1 rangefinder, "
    "an Intel RealSense D435i camera, a Raspberry Pi 4 compute node, and an OpenCR 1.0 controller. "
    "Heavy perception ran on a workstation CPU, communicating via a local Wi-Fi ROS 2 network.\n\n"
    "B. System Integration\n"
    "The manager_node starts perception in TRAINING mode. After collecting 40 frames of the centered target, "
    "the Re-ID gallery is built, and the manager triggers behavior to START.\n\n"
    "C. Experimental Procedure\n"
    "Experiments included: Experiment 1 (Continuous Following), Experiment 2 (Obstacle Avoidance), "
    "Experiment 3 (Target Occlusion), and Experiment 4 (Multiple People).\n\n"
    "D. Experimental Results\n"
    "Real-world trials confirmed stable person following at 0.85 m. To prevent motor voltage drop crashes, "
    "software acceleration limiting restricts linear changes to 0.035 m/s per tick (~0.35 m/s^2) and angular "
    "changes to 0.12 rad/s per tick (~1.20 rad/s^2). The RPLidar subscriber QoS was configured to BEST_EFFORT "
    "to resolve packet drops.\n\n"
    "E. Challenges Encountered\n"
    "1. Lighting Variations: Poor light reduced YOLO confidence and shifted color histogram bins.\n"
    "2. Temporary Target Occlusions: Extended target loss increased recovery time during Re-ID matching.\n"
    "3. Limited Onboard Computing: CPU bottlenecks required executing perception on an external workstation.\n"
    "4. Network Communication: Delays in Wi-Fi transmission introduced minor control latency.\n\n"
    "F. Discussion\n"
    "Real-world testing confirmed the feasibility of the decoupled CPU-only perception-action architecture, "
    "demonstrating stable person following and obstacle avoidance."
)

X_LIMIT = (
    "A. Challenges Encountered\n"
    "1. Lighting Variations: Changes in illumination affected YOLO detection confidence and color histograms.\n"
    "2. Temporary Target Occlusions: Target loss for extended periods increased target reacquisition time.\n"
    "3. Limited Computing Resources: Onboard processor limits required external compute for deep perception.\n"
    "4. Network Communication: Wireless transmission latency introduced minor delays in motor controls.\n"
    "5. Sensor Calibration: Alignment between camera and LiDAR was required for leg exclusion.\n\n"
    "B. System Limitations\n"
    "The system is currently limited to tracking a single target, requires workstation compute, "
    "is restricted to indoor environments, and depends on local-only obstacle costmaps."
)

XI_FUTURE = (
    "A. Multi-Person Tracking: prioritizing targets by face or voice.\n"
    "B. Improved Re-ID: Deployment of OSNet with GPU acceleration.\n"
    "C. ROS 2 Navigation Stack (Nav2): Full costmap integration for global path planning.\n"
    "D. Semantic Mapping and SLAM: Real-time room mapping to recognize doors, shelves, and obstacles.\n"
    "E. Gesture and Voice Recognition: MediaPipe hand signal commands (e.g. Stop, Come Closer).\n"
    "F. Edge AI Deployment: Standalone deployment on embedded NVIDIA Jetson Orin compute nodes.\n"
    "G. Outdoor Navigation: Integrating GPS, IMU, and weather-robust perception."
)

XII_CONCL = (
    "We designed and evaluated an autonomous person-following robot using ROS 2 Humble. "
    "By combining YOLOv8n object detection, ByteTrack tracking, a four-level Re-ID cascade, "
    "and Potential Field obstacle avoidance, we achieved safe, CPU-efficient following "
    "on a TurtleBot3 Waffle Pi. Virtual simulations and real-world trials validate "
    "our system's stability, demonstrating its potential for service robotics in retail, "
    "healthcare, and domestic environments."
)

REFERENCES_LIST = [
    "[1] M. Quigley et al., \"ROS: An open-source Robot Operating System,\" ICRA Workshop on Open Source Software, 2009.",
    "[2] M. Colledanchise and P. Ögren, Behavior trees in robotics and AI: An introduction, CRC Press, 2018.",
    "[3] G. Bradski and A. Kaehler, Learning OpenCV: Computer vision with the OpenCV library, O'Reilly Media, 2008.",
    "[4] J. Redmon et al., \"You only look once: Unified, real-time object detection,\" IEEE CVPR, pp. 779-788, 2016.",
    "[5] N. Wojke, A. Bewley, and D. Pastawski, \"Simple online and realtime tracking with a deep association metric,\" IEEE ICIP, pp. 3645-3649, 2017.",
    "[6] J. S. Esteves et al., \"Physical human-robot interaction based on competency-based model for obesity rehabilitation,\" IEEE Ro-MAN, pp. 681-686, 2012.",
    "[7] A. K. Pandey and R. Alami, \"A framework towards a socially aware mobile robot motion in human environments,\" IROS Workshop on Advances in Service Robotics, 2010."
]

# ─────────────────────────────────────────────────────────────────────────────
# IMAGE VALIDATOR
# ─────────────────────────────────────────────────────────────────────────────
def is_valid_image(path):
    if not path or not os.path.exists(path):
        return False
    try:
        with PIL_Image.open(path) as img:
            img.verify()
        return True
    except Exception:
        return False

# ─────────────────────────────────────────────────────────────────────────────
# DOCX PICTURE ADDER HELPER (SAFE)
# ─────────────────────────────────────────────────────────────────────────────
def add_picture_safe(doc, path, width):
    if is_valid_image(path):
        try:
            doc.add_picture(path, width=width)
            return True
        except Exception as e:
            print(f"[Warning] Failed to add valid image {path} to Docx: {e}")
            
    # Insert a centered text placeholder
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"[IMAGE PLACEHOLDER: {os.path.basename(path)}]")
    r.font.bold = True
    r.font.italic = True
    r.font.color.rgb = RGBColor(120, 120, 120)
    return False

# ─────────────────────────────────────────────────────────────────────────────
# REPORTLAB PICTURE HELPER (SAFE)
# ─────────────────────────────────────────────────────────────────────────────
def get_image_safe(path, width, height):
    if is_valid_image(path):
        try:
            return Image(path, width=width, height=height)
        except Exception as e:
            print(f"[Warning] Failed to add valid image {path} to PDF flowables: {e}")
            
    # Return a spacer or text paragraph as a placeholder
    styles = getSampleStyleSheet()
    err_style = ParagraphStyle(
        'ImgErr', parent=styles['Normal'], fontName='Helvetica-Bold',
        fontSize=10, textColor=colors.HexColor("#777777"), alignment=1
    )
    return Paragraph(f"[IMAGE PLACEHOLDER: {os.path.basename(path)}]", err_style)

# ─────────────────────────────────────────────────────────────────────────────
# WORD DOCUMENT GENERATION (python-docx)
# ─────────────────────────────────────────────────────────────────────────────
def generate_docx():
    print(f"Creating Word document at {DOCX_OUT}...")
    doc = docx.Document()
    
    # Page setup
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    # Styles
    style_normal = doc.styles['Normal']
    font = style_normal.font
    font.name = 'Times New Roman'
    font.size = Pt(10.5)
    font.color.rgb = RGBColor(51, 51, 51)
    
    # Title
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run(TITLE)
    title_run.font.name = 'Times New Roman'
    title_run.font.size = Pt(22)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(15, 76, 129)
    
    # Authors
    author_p = doc.add_paragraph()
    author_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    author_run = author_p.add_run(f"{AUTHORS}\n{INSTITUTION}\n{DATE_STR}")
    author_run.font.name = 'Times New Roman'
    author_run.font.size = Pt(10.5)
    author_run.font.italic = True
    
    doc.add_paragraph().add_run("-" * 80).font.color.rgb = RGBColor(200, 200, 200)
    
    # Abstract
    abs_heading = doc.add_paragraph()
    abs_heading_run = abs_heading.add_run("Abstract")
    abs_heading_run.font.bold = True
    abs_heading_run.font.size = Pt(12)
    abs_heading_run.font.color.rgb = RGBColor(15, 76, 129)
    
    abs_p = doc.add_paragraph()
    abs_run = abs_p.add_run(ABSTRACT_TEXT)
    abs_run.font.italic = True
    abs_run.font.size = Pt(10)
    
    # Keywords
    kw_p = doc.add_paragraph()
    kw_bold = kw_p.add_run("Keywords— ")
    kw_bold.font.bold = True
    kw_bold.font.size = Pt(10)
    kw_text = kw_p.add_run(KEYWORDS_TEXT)
    kw_text.font.size = Pt(10)
    
    doc.add_paragraph().add_run("-" * 80).font.color.rgb = RGBColor(200, 200, 200)
    
    # Section Map
    section_map = [
        ("I. INTRODUCTION", I_INTRO),
        ("II. PROBLEM STATEMENT", II_PROBLEM),
        ("III. PROJECT VISION AND OBJECTIVES", III_VISION),
        ("IV. METHODOLOGY", IV_METHOD),
        ("V. PERCEPTION MODULE: YOLO-BASED PERSON DETECTION SYSTEM", V_PERCEPTION),
        ("VI. TRACKING MODULE: BYTETRACK AND PERSON RE-IDENTIFICATION", VI_TRACKING),
        ("VII. OBSTACLE DETECTION AND AVOIDANCE USING LiDAR", VII_OBSTACLE),
        ("VIII. GAZEBO SIMULATION ENVIRONMENT AND SYSTEM VALIDATION", VIII_SIM),
        ("IX. REAL-WORLD DEPLOYMENT AND EXPERIMENTAL RESULTS", IX_RESULTS),
        ("X. CHALLENGES AND LIMITATIONS", X_LIMIT),
        ("XI. FUTURE WORK", XI_FUTURE),
        ("XII. CONCLUSION", XII_CONCL)
    ]
    
    # Write sections
    for title, text in section_map:
        h = doc.add_paragraph()
        h_run = h.add_run(title)
        h_run.font.name = 'Times New Roman'
        h_run.font.size = Pt(13)
        h_run.font.bold = True
        h_run.font.color.rgb = RGBColor(15, 76, 129)
        
        # Parse content by lines/paragraphs
        lines = text.split("\n\n")
        for line in lines:
            if line.strip().startswith("A. ") or line.strip().startswith("B. ") or line.strip().startswith("C. ") or line.strip().startswith("D. ") or line.strip().startswith("E. ") or line.strip().startswith("F. "):
                # Subsection
                sub_p = doc.add_paragraph()
                sub_run = sub_p.add_run(line.strip())
                sub_run.font.bold = True
                sub_run.font.size = Pt(11)
            else:
                p = doc.add_paragraph()
                p.add_run(line.strip())
                p.paragraph_format.space_after = Pt(8)
                p.paragraph_format.line_spacing = 1.15
                
        # Insert figures
        if "I. INTRODUCTION" in title:
            img_path = os.path.join(IMAGES_DIR, "supermarket_robot_deployment.png")
            add_picture_safe(doc, img_path, width=Inches(5))
            cap = doc.add_paragraph()
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap.add_run("Figure 1: Socially-aware shopping assistant robots deploying in a supermarket retail aisle.").font.italic = True
            
        elif "III. PROJECT VISION" in title:
            img_path = os.path.join(IMAGES_DIR, "presentation1_media", "image6.png")
            add_picture_safe(doc, img_path, width=Inches(4.5))
            cap = doc.add_paragraph()
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap.add_run("Figure 2: TurtleBot3 Waffle Pi mobile robot hardware configuration.").font.italic = True
            
        elif "IV. METHODOLOGY" in title:
            img_path = os.path.join(IMAGES_DIR, "presentation1_media", "image5.png")
            add_picture_safe(doc, img_path, width=Inches(3.5))
            cap = doc.add_paragraph()
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap.add_run("Figure 3: System block diagram / ROS 2 node architecture visualization.").font.italic = True
            
        elif "VIII. GAZEBO SIMULATION" in title:
            img_path = os.path.join(IMAGES_DIR, "presentation1_media", "image10.png")
            add_picture_safe(doc, img_path, width=Inches(4.5))
            cap = doc.add_paragraph()
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap.add_run("Figure 4: Real-time simulation environment validation in Gazebo and Tkinter dashboard telemetry.").font.italic = True

    # Tables section
    doc.add_page_break()
    tb_heading = doc.add_paragraph()
    tb_heading.add_run("SYSTEM DESIGN TABLES & SPECIFICATIONS").font.bold = True
    tb_heading.runs[0].font.size = Pt(13)
    tb_heading.runs[0].font.color.rgb = RGBColor(15, 76, 129)
    
    # Table 1: Safety state definitions
    doc.add_paragraph().add_run("Table I: Safety State Definitions with Hysteresis").font.bold = True
    table1 = doc.add_table(rows=4, cols=4)
    table1.alignment = WD_TABLE_ALIGNMENT.CENTER
    table1.style = 'Light Shading Accent 1'
    hdr_cells = table1.rows[0].cells
    hdr_cells[0].text = 'State'
    hdr_cells[1].text = 'Distance Range (d)'
    hdr_cells[2].text = 'Hysteresis Band'
    hdr_cells[3].text = 'Controller Action'
    
    row_data1 = [
        ('Too Close', '< 0.55 m', '0.10 m', 'Slow reverse / Back up velocity command'),
        ('Safe Follow', '0.55 m - 1.50 m', '0.20 m', 'Maintain target follow distance (0.85 m)'),
        ('Too Far', '> 1.50 m', '0.20 m', 'Proportional acceleration velocity command')
    ]
    for i, data in enumerate(row_data1):
        row_cells = table1.rows[i+1].cells
        for j, val in enumerate(data):
            row_cells[j].text = val
            
    doc.add_paragraph() # Spacer
            
    # Table 2: Technical components
    doc.add_paragraph().add_run("Table II: Technical Components & Project Stack").font.bold = True
    table2 = doc.add_table(rows=6, cols=3)
    table2.alignment = WD_TABLE_ALIGNMENT.CENTER
    table2.style = 'Light Shading Accent 1'
    hdr_cells2 = table2.rows[0].cells
    hdr_cells2[0].text = 'Layer / Component'
    hdr_cells2[1].text = 'Technology / Spec'
    hdr_cells2[2].text = 'System Role'
    
    row_data2 = [
        ('Robot Platform', 'TurtleBot3 Waffle Pi', 'Differential-drive chassis, OpenCR 1.0 hardware compute'),
        ('Onboard Computer', 'Raspberry Pi 4 (4GB)', 'Manages low-level behavior_tree_node & cmd_vel velocity timers'),
        ('Perception Sensors', 'RPLidar A1 + RealSense D435i', '360° laser scan points + RGB-D frames for YOLO processing'),
        ('Software Stack', 'ROS 2 Humble / Ubuntu 22.04', 'System pub/sub topic architecture framework'),
        ('Computer Vision', 'YOLOv8n + ByteTrack + Re-ID', 'Real-time multi-person tracker with appearance biometric locks')
    ]
    for i, data in enumerate(row_data2):
        row_cells = table2.rows[i+1].cells
        for j, val in enumerate(data):
            row_cells[j].text = val
            
    doc.add_paragraph() # Spacer

    # References
    ref_h = doc.add_paragraph()
    ref_h.add_run("References").font.bold = True
    ref_h.runs[0].font.size = Pt(13)
    ref_h.runs[0].font.color.rgb = RGBColor(15, 76, 129)
    
    for ref in REFERENCES_LIST:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.25)
        p.add_run(ref).font.size = Pt(9.5)
        
    doc.save(DOCX_OUT)
    print(f"Word document saved to {DOCX_OUT}.")

# ─────────────────────────────────────────────────────────────────────────────
# PDF GENERATION (reportlab)
# ─────────────────────────────────────────────────────────────────────────────
def generate_pdf():
    print(f"Creating PDF document at {PDF_OUT}...")
    
    doc = SimpleDocTemplate(
        PDF_OUT,
        pagesize=letter,
        rightMargin=54,
        leftMargin=54,
        topMargin=54,
        bottomMargin=54
    )
    
    styles = getSampleStyleSheet()
    
    title_style = ParagraphStyle(
        'DocTitle',
        parent=styles['Normal'],
        fontName='Helvetica-Bold',
        fontSize=18,
        leading=22,
        alignment=1,
        textColor=colors.HexColor("#0F4C81"),
        spaceAfter=15
    )
    
    subtitle_style = ParagraphStyle(
        'DocSubtitle',
        parent=styles['Normal'],
        fontName='Helvetica-Oblique',
        fontSize=10,
        leading=14,
        alignment=1,
        textColor=colors.HexColor("#444444"),
        spaceAfter=15
    )
    
    abstract_heading = ParagraphStyle(
        'AbstractHeading',
        parent=styles['Normal'],
        fontName='Helvetica-Bold',
        fontSize=11,
        leading=15,
        textColor=colors.HexColor("#0F4C81"),
        spaceBefore=10,
        spaceAfter=5
    )
    
    abstract_text = ParagraphStyle(
        'AbstractText',
        parent=styles['Normal'],
        fontName='Helvetica-Oblique',
        fontSize=9,
        leading=13,
        textColor=colors.HexColor("#333333"),
        spaceAfter=15
    )
    
    h1_style = ParagraphStyle(
        'Heading1',
        parent=styles['Normal'],
        fontName='Helvetica-Bold',
        fontSize=13,
        leading=17,
        textColor=colors.HexColor("#0F4C81"),
        spaceBefore=18,
        spaceAfter=8,
        keepWithNext=True
    )
    
    h2_style = ParagraphStyle(
        'Heading2',
        parent=styles['Normal'],
        fontName='Helvetica-Bold',
        fontSize=11,
        leading=15,
        textColor=colors.HexColor("#333333"),
        spaceBefore=12,
        spaceAfter=6,
        keepWithNext=True
    )
    
    body_style = ParagraphStyle(
        'Body',
        parent=styles['Normal'],
        fontName='Helvetica',
        fontSize=9.5,
        leading=13.5,
        spaceAfter=8,
        textColor=colors.HexColor("#333333")
    )
    
    caption_style = ParagraphStyle(
        'Caption',
        parent=styles['Normal'],
        fontName='Helvetica-Oblique',
        fontSize=8,
        leading=11,
        alignment=1,
        textColor=colors.HexColor("#555555"),
        spaceBefore=5,
        spaceAfter=12
    )
    
    table_cell_style = ParagraphStyle(
        'TableCell',
        parent=styles['Normal'],
        fontName='Helvetica',
        fontSize=8.5,
        leading=11,
        textColor=colors.HexColor("#333333")
    )
    
    table_hdr_style = ParagraphStyle(
        'TableHdr',
        parent=styles['Normal'],
        fontName='Helvetica-Bold',
        fontSize=9,
        leading=12,
        textColor=colors.white
    )
    
    ref_style = ParagraphStyle(
        'Reference',
        parent=styles['Normal'],
        fontName='Helvetica',
        fontSize=8.5,
        leading=12,
        leftIndent=15,
        spaceAfter=6,
        textColor=colors.HexColor("#444444")
    )
    
    story = []
    
    story.append(Paragraph(TITLE, title_style))
    story.append(Paragraph(f"{AUTHORS}<br/>{INSTITUTION}<br/>{DATE_STR}", subtitle_style))
    story.append(Spacer(1, 10))
    
    story.append(Paragraph("Abstract", abstract_heading))
    story.append(Paragraph(ABSTRACT_TEXT, abstract_text))
    story.append(Paragraph(f"<b>Keywords—</b> {KEYWORDS_TEXT}", body_style))
    story.append(Spacer(1, 10))
    
    section_map = [
        ("I. INTRODUCTION", I_INTRO),
        ("II. PROBLEM STATEMENT", II_PROBLEM),
        ("III. PROJECT VISION AND OBJECTIVES", III_VISION),
        ("IV. METHODOLOGY", IV_METHOD),
        ("V. PERCEPTION MODULE: YOLO-BASED PERSON DETECTION SYSTEM", V_PERCEPTION),
        ("VI. TRACKING MODULE: BYTETRACK AND PERSON RE-IDENTIFICATION", VI_TRACKING),
        ("VII. OBSTACLE DETECTION AND AVOIDANCE USING LiDAR", VII_OBSTACLE),
        ("VIII. GAZEBO SIMULATION ENVIRONMENT AND SYSTEM VALIDATION", VIII_SIM),
        ("IX. REAL-WORLD DEPLOYMENT AND EXPERIMENTAL RESULTS", IX_RESULTS),
        ("X. CHALLENGES AND LIMITATIONS", X_LIMIT),
        ("XI. FUTURE WORK", XI_FUTURE),
        ("XII. CONCLUSION", XII_CONCL)
    ]
    
    for title, text in section_map:
        story.append(Paragraph(title, h1_style))
        lines = text.split("\n\n")
        for line in lines:
            if line.strip().startswith("A. ") or line.strip().startswith("B. ") or line.strip().startswith("C. ") or line.strip().startswith("D. ") or line.strip().startswith("E. ") or line.strip().startswith("F. "):
                story.append(Paragraph(line.strip(), h2_style))
            else:
                story.append(Paragraph(line.strip(), body_style))
                
        if "I. INTRODUCTION" in title:
            img_path = os.path.join(IMAGES_DIR, "supermarket_robot_deployment.png")
            story.append(KeepTogether([
                get_image_safe(img_path, width=420, height=230),
                Paragraph("Figure 1: Socially-aware shopping assistant robots deploying in a supermarket retail aisle.", caption_style)
            ]))
                
        elif "III. PROJECT VISION" in title:
            img_path = os.path.join(IMAGES_DIR, "presentation1_media", "image6.png")
            story.append(KeepTogether([
                get_image_safe(img_path, width=320, height=180),
                Paragraph("Figure 2: TurtleBot3 Waffle Pi mobile robot hardware configuration.", caption_style)
            ]))
                
        elif "IV. METHODOLOGY" in title:
            img_path = os.path.join(IMAGES_DIR, "presentation1_media", "image5.png")
            story.append(KeepTogether([
                get_image_safe(img_path, width=320, height=180),
                Paragraph("Figure 3: System block diagram / ROS 2 node architecture visualization.", caption_style)
            ]))
                
        elif "VIII. GAZEBO SIMULATION" in title:
            img_path = os.path.join(IMAGES_DIR, "presentation1_media", "image10.png")
            story.append(KeepTogether([
                get_image_safe(img_path, width=380, height=220),
                Paragraph("Figure 4: Real-time simulation environment validation in Gazebo and Tkinter dashboard telemetry.", caption_style)
            ]))

    story.append(PageBreak())
    
    story.append(Paragraph("SYSTEM DESIGN TABLES & SPECIFICATIONS", h1_style))
    story.append(Spacer(1, 10))
    
    story.append(Paragraph("Table I: Safety State Definitions with Hysteresis", h2_style))
    t1_headers = [Paragraph("State", table_hdr_style), Paragraph("Distance Range (d)", table_hdr_style), Paragraph("Hysteresis Band", table_hdr_style), Paragraph("Controller Action", table_hdr_style)]
    t1_rows = [
        [Paragraph("Too Close", table_cell_style), Paragraph("< 0.55 m", table_cell_style), Paragraph("0.10 m", table_cell_style), Paragraph("Slow reverse / Back up velocity command", table_cell_style)],
        [Paragraph("Safe Follow", table_cell_style), Paragraph("0.55 m - 1.50 m", table_cell_style), Paragraph("0.20 m", table_cell_style), Paragraph("Maintain target follow distance (0.85 m)", table_cell_style)],
        [Paragraph("Too Far", table_cell_style), Paragraph("> 1.50 m", table_cell_style), Paragraph("0.20 m", table_cell_style), Paragraph("Proportional acceleration velocity command", table_cell_style)]
    ]
    t1_data = [t1_headers] + t1_rows
    t1 = Table(t1_data, colWidths=[100, 110, 100, 190])
    t1.setStyle(TableStyle([
        ('BACKGROUND', (0,0), (-1,0), colors.HexColor("#0F4C81")),
        ('GRID', (0,0), (-1,-1), 0.5, colors.HexColor("#CCCCCC")),
        ('ROWBACKGROUNDS', (0,1), (-1,-1), [colors.white, colors.HexColor("#F9F9F9")]),
        ('TOPPADDING', (0,0), (-1,-1), 6),
        ('BOTTOMPADDING', (0,0), (-1,-1), 6),
    ]))
    story.append(t1)
    story.append(Spacer(1, 20))
    
    story.append(Paragraph("Table II: Technical Components & Project Stack", h2_style))
    t2_headers = [Paragraph("Layer / Component", table_hdr_style), Paragraph("Technology / Spec", table_hdr_style), Paragraph("System Role", table_hdr_style)]
    t2_rows = [
        [Paragraph("Robot Platform", table_cell_style), Paragraph("TurtleBot3 Waffle Pi", table_cell_style), Paragraph("Differential-drive chassis, OpenCR 1.0 hardware compute", table_cell_style)],
        [Paragraph("Onboard Computer", table_cell_style), Paragraph("Raspberry Pi 4 (4GB)", table_cell_style), Paragraph("Manages low-level behavior_tree_node & cmd_vel velocity timers", table_cell_style)],
        [Paragraph("Perception Sensors", table_cell_style), Paragraph("RPLidar A1 + RealSense D435i", table_cell_style), Paragraph("360° laser scan points + RGB-D frames for YOLO processing", table_cell_style)],
        [Paragraph("Software Stack", table_cell_style), Paragraph("ROS 2 Humble / Ubuntu 22.04", table_cell_style), Paragraph("System pub/sub topic architecture framework", table_cell_style)],
        [Paragraph("Computer Vision", table_cell_style), Paragraph("YOLOv8n + ByteTrack + Re-ID", table_cell_style), Paragraph("Real-time multi-person tracker with appearance biometric locks", table_cell_style)]
    ]
    t2_data = [t2_headers] + t2_rows
    t2 = Table(t2_data, colWidths=[120, 150, 230])
    t2.setStyle(TableStyle([
        ('BACKGROUND', (0,0), (-1,0), colors.HexColor("#0F4C81")),
        ('GRID', (0,0), (-1,-1), 0.5, colors.HexColor("#CCCCCC")),
        ('ROWBACKGROUNDS', (0,1), (-1,-1), [colors.white, colors.HexColor("#F9F9F9")]),
        ('TOPPADDING', (0,0), (-1,-1), 6),
        ('BOTTOMPADDING', (0,0), (-1,-1), 6),
    ]))
    story.append(t2)
    story.append(Spacer(1, 20))

    story.append(Paragraph("References", h1_style))
    for ref in REFERENCES_LIST:
        story.append(Paragraph(ref, ref_style))
        
    doc.build(story)
    print(f"PDF document saved to {PDF_OUT}.")

# ─────────────────────────────────────────────────────────────────────────────
# MARKDOWN GENERATION
# ─────────────────────────────────────────────────────────────────────────────
def generate_md():
    print(f"Creating Markdown document at {MD_OUT}...")
    
    content = []
    content.append(f"# {TITLE}\n")
    content.append(f"**Authors:** {AUTHORS}  ")
    content.append(f"**Institution:** {INSTITUTION}  ")
    content.append(f"**Date:** {DATE_STR}  \n")
    content.append("---")
    content.append(f"## Abstract\n{ABSTRACT_TEXT}\n")
    content.append(f"**Keywords—** {KEYWORDS_TEXT}\n")
    content.append("---")
    
    section_map = [
        ("I. INTRODUCTION", I_INTRO),
        ("II. PROBLEM STATEMENT", II_PROBLEM),
        ("III. PROJECT VISION AND OBJECTIVES", III_VISION),
        ("IV. METHODOLOGY", IV_METHOD),
        ("V. PERCEPTION MODULE: YOLO-BASED PERSON DETECTION SYSTEM", V_PERCEPTION),
        ("VI. TRACKING MODULE: BYTETRACK AND PERSON RE-IDENTIFICATION", VI_TRACKING),
        ("VII. OBSTACLE DETECTION AND AVOIDANCE USING LiDAR", VII_OBSTACLE),
        ("VIII. GAZEBO SIMULATION ENVIRONMENT AND SYSTEM VALIDATION", VIII_SIM),
        ("IX. REAL-WORLD DEPLOYMENT AND EXPERIMENTAL RESULTS", IX_RESULTS),
        ("X. CHALLENGES AND LIMITATIONS", X_LIMIT),
        ("XI. FUTURE WORK", XI_FUTURE),
        ("XII. CONCLUSION", XII_CONCL)
    ]
    
    for title, text in section_map:
        content.append(f"## {title}\n")
        lines = text.split("\n\n")
        for line in lines:
            if line.strip().startswith("A. ") or line.strip().startswith("B. ") or line.strip().startswith("C. ") or line.strip().startswith("D. ") or line.strip().startswith("E. ") or line.strip().startswith("F. "):
                content.append(f"### {line.strip()}\n")
            else:
                content.append(f"{line.strip()}\n")
                
        # Insert image references
        if "I. INTRODUCTION" in title:
            content.append("\n![Figure 1: Socially-aware shopping assistant robots deploying in a supermarket retail aisle](../IMAGES/supermarket_robot_deployment.png)\n")
        elif "III. PROJECT VISION" in title:
            content.append("\n![Figure 2: TurtleBot3 Waffle Pi mobile robot hardware configuration](../IMAGES/presentation1_media/image6.png)\n")
        elif "IV. METHODOLOGY" in title:
            content.append("\n![Figure 3: System block diagram / ROS 2 node architecture visualization](../IMAGES/presentation1_media/image5.png)\n")
        elif "VIII. GAZEBO SIMULATION" in title:
            content.append("\n![Figure 4: Real-time simulation environment validation in Gazebo and Tkinter dashboard telemetry](../IMAGES/presentation1_media/image10.png)\n")
            
    content.append("## SYSTEM DESIGN TABLES & SPECIFICATIONS\n")
    content.append("### Table I: Safety State Definitions with Hysteresis\n")
    content.append("| State | Distance Range (d) | Hysteresis Band | Controller Action |")
    content.append("| :--- | :--- | :--- | :--- |")
    content.append("| **Too Close** | < 0.55 m | 0.10 m | Slow reverse / Back up velocity command |")
    content.append("| **Safe Follow** | 0.55 m - 1.50 m | 0.20 m | Maintain target follow distance (0.85 m) |")
    content.append("| **Too Far** | > 1.50 m | 0.20 m | Proportional acceleration velocity command |\n")
    
    content.append("### Table II: Technical Components & Project Stack\n")
    content.append("| Layer / Component | Technology / Spec | System Role |")
    content.append("| :--- | :--- | :--- |")
    content.append("| Robot Platform | TurtleBot3 Waffle Pi | Differential-drive chassis, OpenCR 1.0 hardware compute |")
    content.append("| Onboard Computer | Raspberry Pi 4 (4GB) | Manages low-level behavior_tree_node & cmd_vel velocity timers |")
    content.append("| Perception Sensors | RPLidar A1 + RealSense D435i | 360° laser scan points + RGB-D frames for YOLO processing |")
    content.append("| Software Stack | ROS 2 Humble / Ubuntu 22.04 | System pub/sub topic architecture framework |")
    content.append("| Computer Vision | YOLOv8n + ByteTrack + Re-ID | Real-time multi-person tracker with appearance biometric locks |\n")
    
    content.append("## References\n")
    for ref in REFERENCES_LIST:
        content.append(f"* {ref}")
        
    with open(MD_OUT, "w") as f:
        f.write("\n".join(content))
    print(f"Markdown document saved to {MD_OUT}.")

# ─────────────────────────────────────────────────────────────────────────────
# DYNAMIC LOAD OF AGY BEST REPORT CONTENT
# ─────────────────────────────────────────────────────────────────────────────
def load_agy_best_text():
    agy_path = "/home/ganeshna/person_follower_robot_project/AGY BEST Autonomous_Person_Following_Robot_Complete_IEEE_Paper.md"
    if not os.path.exists(agy_path):
        print(f"Warning: {agy_path} not found. Using default script text variables.")
        return
    
    print(f"Loading and parsing AGY best report from {agy_path}...")
    with open(agy_path, "r") as f:
        text = f.read()
    
    # Extract Abstract
    abstract_start = text.find("***Abstract*—")
    keywords_start = text.find("***Keywords*—")
    if abstract_start != -1 and keywords_start != -1:
        global ABSTRACT_TEXT
        ABSTRACT_TEXT = text[abstract_start + len("***Abstract*—"):keywords_start].strip()
        ABSTRACT_TEXT = ABSTRACT_TEXT.replace("**", "")

    # Extract sections
    sections = [
        ("## I. Introduction", "## II. Problem Statement", "I_INTRO"),
        ("## II. Problem Statement", "## III. Project Vision and Objectives", "II_PROBLEM"),
        ("## III. Project Vision and Objectives", "## IV. Methodology", "III_VISION"),
        ("## IV. Methodology", "## V. Perception Module", "IV_METHOD"),
        ("## V. Perception Module: YOLO-Based Person Detection System", "## VI. Tracking Module", "V_PERCEPTION"),
        ("## VI. Tracking Module: ByteTrack and Person Re-Identification", "## VII. Obstacle Detection", "VI_TRACKING"),
        ("## VII. Obstacle Detection and Avoidance Using LiDAR", "## VIII. Gazebo Simulation", "VII_OBSTACLE"),
        ("## VIII. Gazebo Simulation Environment and System Validation", "## IX. Real-World Deployment", "VIII_SIM"),
        ("## IX. Real-World Deployment and Experimental Results", "## X. Challenges and Limitations", "IX_RESULTS"),
        ("## X. Challenges and Limitations", "## XI. Future Work", "X_LIMIT"),
        ("## XI. Future Work", "## XII. Conclusion", "XI_FUTURE"),
        ("## XII. Conclusion", "## Acknowledgment", "XII_CONCL")
    ]
    
    for start_tag, end_tag, var_name in sections:
        start_idx = text.find(start_tag)
        end_idx = text.find(end_tag)
        if start_idx != -1 and end_idx != -1:
            sec_text = text[start_idx + len(start_tag):end_idx].strip()
            # Apply range gating check to LiDAR avoidance
            if var_name == "VII_OBSTACLE":
                sec_text = sec_text.replace(
                    "Rays within a \\pm 0.38 rad window around the target's angle are excluded to prevent the target from being detected as an obstacle.",
                    "Rays within a \\pm 0.38 rad window around the target's bearing angle are excluded to prevent the target's legs from being detected as an obstacle. To prevent ignoring static obstacles like shelves located in the target's direction, a range gating condition (r >= 0.45 m and abs(r - d_target) < 0.30 m) is applied, ensuring that any close-range object or background shelf is treated as an active obstacle."
                )
            
            # Clean heading markdown syntax (### A. -> A.)
            sec_lines = []
            for line in sec_text.split("\n"):
                if line.strip().startswith("### "):
                    sec_lines.append(line.strip()[4:])
                else:
                    sec_lines.append(line)
            globals()[var_name] = "\n\n".join(sec_lines)

# ─────────────────────────────────────────────────────────────────────────────
# MAIN RUNNER
# ─────────────────────────────────────────────────────────────────────────────
def main():
    load_agy_best_text()
    generate_docx()
    generate_pdf()
    generate_md()

if __name__ == '__main__':
    main()
